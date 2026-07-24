"""
parser.py
---------
Responsible for extracting raw text out of uploaded resume files
(PDF or DOCX). This is step 1 of the pipeline: file -> plain text.
Step 2 (text -> structured fields) lives in extractor.py.
"""
import os
import pdfplumber
from PyPDF2 import PdfReader
from docx import Document


class UnsupportedFileTypeError(Exception):
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """
    Try pdfplumber first (better layout handling), fall back to PyPDF2
    if pdfplumber fails or returns nothing (e.g. certain malformed PDFs).
    """
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
    except Exception:
        text = ""

    if not text.strip():
        # Fallback to PyPDF2
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        except Exception:
            pass

    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a .docx file, including tables."""
    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of any tables (common in resumes for skills sections)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n".join(parts).strip()


def extract_text(file_path: str) -> str:
    """
    Entry point: detects file type by extension and routes to the
    correct extractor. Raises UnsupportedFileTypeError for anything else.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Only .pdf and .docx are supported."
        )
